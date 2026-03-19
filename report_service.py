import os
import re
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.drawing.image import Image as XlImage
from openpyxl.styles import Font, PatternFill, Alignment
from db_config import get_db
from datetime import datetime

class DisclosureReportService:
    @staticmethod
    def generate_report(company_id, year):
        template_path = os.path.join(os.path.dirname(__file__), 'documents', 'disclosure_template.docx')
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")

        doc = Document(template_path)
        
        with get_db() as conn:
            # 1. 회사 정보 조회
            company = conn.execute('SELECT * FROM isd_companies WHERE id=?', (company_id,)).fetchone()
            # 2. 답변 조회
            answers_rows = conn.execute(
                'SELECT question_id, value FROM isd_answers WHERE company_id=? AND year=? AND deleted_at IS NULL',
                (company_id, year)
            ).fetchall()
            answers = {r['question_id']: r['value'] for r in answers_rows}

        # 3. 데이터 추출 및 가공
        def get_val(qid, default='0'):
            v = answers.get(qid, default)
            if v is None or v == '': return default
            return v

        def set_cell_right(cell, text):
            """셀 텍스트 설정 후 오른쪽 정렬"""
            cell.text = text
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        def format_num(v):
            try:
                if isinstance(v, str):
                    v = v.replace(',', '')
                return f"{float(v):,.0f}"
            except (ValueError, TypeError):
                return str(v)

        # 투자액 계산
        it_a = get_val('Q2')
        sec_b1 = get_val('Q4')
        sec_b2 = get_val('Q5')
        sec_b3 = get_val('Q6')
        
        try:
            total_b = float(str(sec_b1).replace(',', '') or 0) + \
                      float(str(sec_b2).replace(',', '') or 0) + \
                      float(str(sec_b3).replace(',', '') or 0)
        except ValueError:
            total_b = 0

        ratio_inv = "0"
        try:
            val_a = float(str(it_a).replace(',', '') or 0)
            if val_a > 0:
                ratio_inv = f"{(total_b / val_a) * 100:.2f}"
        except ValueError:
            pass

        # 인력 계산
        total_emp = get_val('Q10')
        it_emp = get_val('Q28')
        internal = get_val('Q11')
        external = get_val('Q12')
        try:
            total_d = float(str(internal).replace(',', '') or 0) + \
                      float(str(external).replace(',', '') or 0)
        except ValueError:
            total_d = 0
            
        ratio_per = "0"
        try:
            val_c = float(str(it_emp).replace(',', '') or 0)
            if val_c > 0:
                ratio_per = f"{(total_d / val_c) * 100:.2f}"
        except ValueError:
            pass

        # 4. 워드 문서 필드 채우기
        # doc.tables[0]은 전체 컨테이너 테이블
        container = doc.tables[0]
        main_cell = container.rows[0].cells[0]
        nested_tables = main_cell.tables
        
        # 작성 기준일 채우기 (공시 연도의 마지막 날: YYYY년 12월 31일)
        date_ymd = f"{year % 100:02d}. 12. 31."
        for para in main_cell.paragraphs:
            if '작성 기준일' in para.text:
                runs = para.runs
                if runs:
                    full = ''.join(r.text or '' for r in runs)
                    new_full = re.sub(r'20\s*\.\s*\.', date_ymd, full)
                    if new_full != full:
                        runs[0].text = new_full
                        for r in runs[1:]:
                            r.text = ''
                break

        if len(nested_tables) >= 2:
            # Table 1 (Title)
            nested_tables[0].rows[0].cells[0].text = f"({company['name']}) 정보보호 현황"

            # Table 2 (Main)
            t2 = nested_tables[1]
            set_cell_right(t2.rows[0].cells[3], f"{format_num(it_a)} 원")
            set_cell_right(t2.rows[1].cells[3], f"{format_num(total_b)} 원")

            # 주요 투자 항목 (Q27)
            if 'Q27' in answers:
                try:
                    items = json.loads(answers['Q27'])
                    item_text = "\n".join([f"- {i.get('항목명', '')}: {format_num(i.get('금액', 0))}원" for i in items if i.get('항목명')])
                    t2.rows[3].cells[3].text = item_text
                except:
                    t2.rows[3].cells[3].text = str(answers['Q27'])

            set_cell_right(t2.rows[4].cells[3], f"{ratio_inv} %")

            set_cell_right(t2.rows[6].cells[3], f"{format_num(total_emp)} 명")
            set_cell_right(t2.rows[7].cells[3], f"{format_num(it_emp)} 명")
            set_cell_right(t2.rows[8].cells[3], f"{format_num(internal)} 명")
            set_cell_right(t2.rows[9].cells[3], f"{format_num(external)} 명")
            set_cell_right(t2.rows[10].cells[3], f"{format_num(total_d)} 명")
            set_cell_right(t2.rows[11].cells[3], f"{ratio_per} %")
            
            # CISO/CPO (Nested Table in Row 12, Cell 3 of T2)
            sub_tables = t2.rows[12].cells[3].tables
            if sub_tables:
                target_tbl = sub_tables[0]
                if answers.get('Q13') not in ('YES', 'Y'):
                    # Q13 미지정: 데이터 행 첫 번째 셀에 N/A 표시
                    for r_idx in [1, 2]:
                        if r_idx < len(target_tbl.rows):
                            for cell in target_tbl.rows[r_idx].cells:
                                cell.text = 'N/A' if cell == target_tbl.rows[r_idx].cells[1] else ''
                elif 'Q14' in answers:
                    try:
                        ciso_data_word = json.loads(answers['Q14'])

                        # Q29 주요 활동 내역 → person별 집계
                        q29_by_person = {'CISO': [], 'CPO': []}
                        if 'Q29' in answers:
                            q29_rows = json.loads(answers['Q29'])
                            for r in q29_rows:
                                person = r.get('person', '')
                                if person in q29_by_person and r.get('activity_type'):
                                    count = r.get('count', '')
                                    line = f"- {r['activity_type']}: {r.get('detail', '')} ({count}회)" if count else f"- {r['activity_type']}: {r.get('detail', '')}"
                                    q29_by_person[person].append(line)

                        for row in ciso_data_word:
                            person_type = row.get('type', '')
                            r_idx = 1 if person_type == 'CISO' else (2 if person_type == 'CPO' else None)
                            if r_idx and r_idx < len(target_tbl.rows):
                                target_tbl.rows[r_idx].cells[1].text = row.get('position', '')
                                target_tbl.rows[r_idx].cells[2].text = row.get('is_officer', '')
                                target_tbl.rows[r_idx].cells[3].text = row.get('is_concurrent', '')
                                target_tbl.rows[r_idx].cells[4].text = '\n'.join(q29_by_person.get(person_type, []))
                    except:
                        pass
            
            # 3. 인증 현황 (Q16)
            if 'Q16' in answers:
                try:
                    certs = json.loads(answers['Q16'])
                    cert_text = "\n".join([f"- {i.get('인증명', '')} ({i.get('기관', '')}, {i.get('유효기간', '')})" for i in certs if i.get('인증명')])
                    t2.rows[14].cells[3].text = cert_text
                except:
                    t2.rows[14].cells[3].text = str(answers['Q16'])
            
            # 4. 활동 현황 (Category 4) - Q17이 YES인 경우에만 하위 항목 표시
            with get_db() as conn:
                act_qs = conn.execute('SELECT id, text FROM isd_questions WHERE category_id=4 AND type="yes_no"').fetchall()
                acts = []
                if answers.get('Q17') in ('YES', 'Y'):
                    for q in act_qs:
                        if q['id'] == 'Q17':
                            continue
                        if answers.get(q['id']) in ('YES', 'Y'):
                            q_text = q['text'].replace(' 있나요?', '').replace(' 하고 있나요?', '')
                            acts.append(f"- {q_text}")
                t2.rows[15].cells[3].text = "\n".join(acts)

        # Footer Confirmation (Container Row 1)
        container.rows[1].cells[0].text = f"{company['name']} 대표이사 OOO는 상기 공시 내용에 거짓이 없음을 확인하였습니다.\n\n{datetime.now().strftime('%Y. %m. %d.')}\n\n(회사 대표이사 직인)"

        output_filename = f"정보보호공시_{company['name']}_{year}.docx"
        output_dir = os.path.join(os.path.dirname(__file__), 'uploads')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        return output_path, output_filename

    @staticmethod
    def generate_excel_report(company_id, year):
        """엑셀 공시 양식에 DB 데이터를 채워 반환"""
        template_path = os.path.join(os.path.dirname(__file__), 'documents', 'disclosure_template.xlsx')
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Excel template not found: {template_path}")

        with get_db() as conn:
            company = conn.execute('SELECT * FROM isd_companies WHERE id=?', (company_id,)).fetchone()
            answers_rows = conn.execute(
                'SELECT question_id, value FROM isd_answers WHERE company_id=? AND year=? AND deleted_at IS NULL',
                (company_id, year)
            ).fetchall()
            answers = {r['question_id']: r['value'] for r in answers_rows}
            act_qs = conn.execute(
                'SELECT id, text FROM isd_questions WHERE category_id=4 AND type="yes_no"'
            ).fetchall()

        def get_val(qid, default='0'):
            v = answers.get(qid, default)
            return default if (v is None or v == '') else v

        def format_num(v):
            try:
                if isinstance(v, str):
                    v = v.replace(',', '')
                return f"{float(v):,.0f}"
            except (ValueError, TypeError):
                return str(v)

        # 투자액 계산
        it_a = get_val('Q2')
        sec_b1 = get_val('Q4')
        sec_b2 = get_val('Q5')
        sec_b3 = get_val('Q6')
        try:
            total_b = (float(str(sec_b1).replace(',', '') or 0) +
                       float(str(sec_b2).replace(',', '') or 0) +
                       float(str(sec_b3).replace(',', '') or 0))
        except ValueError:
            total_b = 0

        ratio_inv = "0"
        try:
            val_a = float(str(it_a).replace(',', '') or 0)
            if val_a > 0:
                ratio_inv = f"{(total_b / val_a) * 100:.2f}"
        except ValueError:
            pass

        # 인력 계산
        total_emp = get_val('Q10')
        it_emp = get_val('Q28')
        internal = get_val('Q11')
        external = get_val('Q12')
        try:
            total_d = (float(str(internal).replace(',', '') or 0) +
                       float(str(external).replace(',', '') or 0))
        except ValueError:
            total_d = 0

        ratio_per = "0"
        try:
            val_c = float(str(it_emp).replace(',', '') or 0)
            if val_c > 0:
                ratio_per = f"{(total_d / val_c) * 100:.2f}"
        except ValueError:
            pass

        # 주요 투자 항목 (Q27)
        q27_text = ''
        if 'Q27' in answers:
            try:
                items = json.loads(answers['Q27'])
                q27_text = '\n'.join([
                    f"- {i.get('항목명', '')}: {format_num(i.get('금액', 0))}원"
                    for i in items if i.get('항목명')
                ])
            except Exception:
                q27_text = str(answers['Q27'])

        # 인증 현황 (Q16)
        q16_text = ''
        if 'Q16' in answers:
            try:
                certs = json.loads(answers['Q16'])
                q16_text = '\n'.join([
                    f"- {i.get('인증명', '')} ({i.get('기관', '')}, {i.get('유효기간', '')})"
                    for i in certs if i.get('인증명')
                ])
            except Exception:
                q16_text = str(answers['Q16'])

        # 활동 현황 (category 4 yes_no) - Q17이 YES인 경우에만 하위 항목 표시
        acts = []
        if answers.get('Q17') in ('YES', 'Y'):
            for q in act_qs:
                if q['id'] == 'Q17':
                    continue
                if answers.get(q['id']) in ('YES', 'Y'):
                    q_text = q['text'].replace(' 있나요?', '').replace(' 하고 있나요?', '')
                    acts.append(f"- {q_text}")
        act_text = '\n'.join(acts)

        # CISO/CPO 데이터 (Sheet 2에만 기입)
        ciso_data = []
        if 'Q14' in answers:
            try:
                ciso_data = json.loads(answers['Q14'])
            except Exception:
                pass

        # Q29 주요 활동 내역 → person별 집계
        q29_by_person = {'CISO': [], 'CPO': []}
        if 'Q29' in answers:
            try:
                q29_rows = json.loads(answers['Q29'])
                for r in q29_rows:
                    person = r.get('person', '')
                    if person in q29_by_person and r.get('activity_type'):
                        count = r.get('count', '')
                        line = f"- {r['activity_type']}: {r.get('detail', '')} ({count}회)" if count else f"- {r['activity_type']}: {r.get('detail', '')}"
                        q29_by_person[person].append(line)
            except Exception:
                pass

        wb = openpyxl.load_workbook(template_path)

        # Sheet 1: 정보보호 공시양식
        ws1 = wb['정보보호 공시양식']
        ws1['A1'] = f"({company['name']}) 정보보호 현황"
        ws1['D3'] = f"{format_num(it_a)} 원"
        ws1['D4'] = f"{format_num(total_b)} 원"
        ws1['D5'] = q27_text
        ws1['D7'] = f"{ratio_inv} %"
        ws1['D9'] = f"{format_num(total_emp)} 명"
        ws1['D10'] = f"{format_num(it_emp)} 명"
        ws1['D11'] = f"{format_num(internal)} 명"
        ws1['D12'] = f"{format_num(external)} 명"
        ws1['D13'] = f"{format_num(total_d)} 명"
        ws1['D14'] = f"{ratio_per} %"

        # CISO/CPO 지정 현황 (D15) - Q13 지정 여부 + Q14 상세 요약
        if answers.get('Q13') not in ('YES', 'Y'):
            ws1['D15'] = 'N/A'
        else:
            ciso_lines = []
            for row in ciso_data:
                person_type = row.get('type', '')
                if not person_type:
                    continue
                position = row.get('position', '')
                is_officer = row.get('is_officer', '')
                is_concurrent = row.get('is_concurrent', '')
                parts = [p for p in [position, is_officer, is_concurrent] if p]
                line = f"{person_type}: 지정 ({', '.join(parts)})" if parts else f"{person_type}: 지정"
                ciso_lines.append(line)
            ws1['D15'] = '\n'.join(ciso_lines) if ciso_lines else '지정'

        ws1['D17'] = q16_text
        ws1['D18'] = act_text

        # Sheet 2: CISO_CPO 현황
        ws2 = wb['CISO_CPO 현황']
        for row in ciso_data:
            person_type = row.get('type', '')
            r_idx = 3 if person_type == 'CISO' else (4 if person_type == 'CPO' else None)
            if r_idx:
                ws2.cell(row=r_idx, column=2).value = row.get('position', '')
                ws2.cell(row=r_idx, column=3).value = row.get('is_officer', '')
                ws2.cell(row=r_idx, column=4).value = row.get('is_concurrent', '')
                ws2.cell(row=r_idx, column=5).value = '\n'.join(q29_by_person.get(person_type, []))

        # Sheet 3~: 질문별 증빙 파일 시트
        upload_base = os.path.join(os.path.dirname(__file__), 'uploads', 'disclosure', company_id, str(year))
        with get_db() as conn:
            ev_rows = conn.execute('''
                SELECT e.question_id, e.file_name, e.file_type, e.file_size, e.uploaded_at, e.file_url,
                       q.display_number, q.text
                FROM isd_evidence e
                JOIN isd_questions q ON e.question_id = q.id
                WHERE e.company_id=? AND e.year=?
                ORDER BY q.sort_order, e.uploaded_at
            ''', (company_id, year)).fetchall()

        # 질문별로 그룹핑
        from collections import OrderedDict
        ev_by_q = OrderedDict()
        for ev in ev_rows:
            qid = ev['question_id']
            if qid not in ev_by_q:
                ev_by_q[qid] = {'display_number': ev['display_number'], 'text': ev['text'], 'files': []}
            ev_by_q[qid]['files'].append(ev)

        IMAGE_EXTS = {'jpg', 'jpeg', 'png', 'gif', 'bmp'}
        header_fill = PatternFill('solid', fgColor='D9E1F2')
        header_font = Font(bold=True)

        for qid, qdata in ev_by_q.items():
            disp = qdata['display_number'] or qid
            q_text = qdata['text'] or ''
            # 시트명: 31자 제한
            sheet_name = f"{disp} 증빙"
            if len(sheet_name) > 31:
                sheet_name = sheet_name[:31]
            ws = wb.create_sheet(title=sheet_name)

            # 헤더
            ws['A1'] = f"[{disp}] {q_text}"
            ws['A1'].font = Font(bold=True, size=11)
            ws.merge_cells('A1:D1')

            headers = ['파일명', '유형', '크기(KB)', '업로드 일시']
            for ci, h in enumerate(headers, 1):
                cell = ws.cell(row=2, column=ci, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')

            ws.column_dimensions['A'].width = 35
            ws.column_dimensions['B'].width = 10
            ws.column_dimensions['C'].width = 12
            ws.column_dimensions['D'].width = 22

            data_row = 3
            for ev in qdata['files']:
                # 파일 실제 경로
                url_filename = ev['file_url'].rsplit('/', 1)[-1]
                file_path = os.path.join(upload_base, url_filename)

                ws.cell(row=data_row, column=1, value=ev['file_name'])
                ws.cell(row=data_row, column=2, value=ev['file_type'])
                ws.cell(row=data_row, column=3, value=round((ev['file_size'] or 0) / 1024, 1))
                ws.cell(row=data_row, column=4, value=str(ev['uploaded_at'])[:19])
                data_row += 1

                # 이미지 파일은 시트에 삽입
                if ev['file_type'] in IMAGE_EXTS and os.path.exists(file_path):
                    try:
                        img = XlImage(file_path)
                        # 최대 너비 400px 기준 비율 축소
                        max_px = 400
                        if img.width > max_px:
                            ratio = max_px / img.width
                            img.width = int(img.width * ratio)
                            img.height = int(img.height * ratio)
                        ws.add_image(img, f'A{data_row}')
                        # 이미지 높이만큼 행 높이 확보 (1px ≈ 0.75pt)
                        ws.row_dimensions[data_row].height = img.height * 0.75
                        data_row += 1
                    except Exception:
                        pass

        output_filename = f"정보보호공시_{company['name']}_{year}.xlsx"
        output_dir = os.path.join(os.path.dirname(__file__), 'uploads')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
        wb.save(output_path)
        return output_path, output_filename
